"""
KCI 저자 식별 알고리즘 통합 전략보고서 (Final) DOCX 생성기
- 수학 알고리즘 + LLM 비교 분석 + Biblo 온톨로지 전략 적용
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ── Style setup ──
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

C_NAVY = RGBColor(0x1A, 0x1A, 0x2E)
C_BLUE = RGBColor(0x0F, 0x4C, 0x81)
C_GREEN = RGBColor(0x1B, 0x7A, 0x4D)
C_PURPLE = RGBColor(0x6C, 0x3E, 0xB6)
C_ORANGE = RGBColor(0xE6, 0x51, 0x00)
C_RED = RGBColor(0xC0, 0x39, 0x2B)
C_GRAY = RGBColor(0x66, 0x66, 0x66)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_TEAL = RGBColor(0x00, 0x89, 0x7B)


def add_heading(text, level=1, color=C_NAVY):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return h


def add_para(text, bold=False, color=None, size=10, align=None, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(text, level=0, bold_prefix="", color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Pt(18 + level * 18)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        if color:
            run.font.color.rgb = color
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return p


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def make_table(headers, rows, col_widths=None, header_color="0F4C81"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = C_WHITE
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        set_cell_shading(cell, header_color)

    # Rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            if r % 2 == 1:
                set_cell_shading(cell, "F5F5FA")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table


def add_code_block(text, font_size=9):
    """Add a monospaced code block."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Consolas'
    return p


# ════════════════════════════════════════
#  표지
# ════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

add_para("RIMS 데이터 품질 고도화 프로젝트", size=12, color=C_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para("KCI 논문 저자 식별 알고리즘 통합 전략 보고서 (Final)", size=22, bold=True, color=C_NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("저자 역할 분류 + 동명이인 구분 + 동일저자 추적", size=14, color=C_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("+ LLM 알고리즘 비교 + 온톨로지 데이터 구조", size=14, color=C_TEAL, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(6):
    doc.add_paragraph()

add_para("성균관대학교 | 연구성과관리시스템(RIMS)", size=11, color=C_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("2026. 02", size=11, color=C_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ════════════════════════════════════════
#  목차
# ════════════════════════════════════════
add_heading("목차", level=1, color=C_NAVY)

toc_items = [
    ("Part I: 현황 분석", True),
    ("  1. RIMS 데이터 현황", False),
    ("  2. TPI_DVS_CD 분포", False),
    ("  3. 동명이인 현황", False),
    ("  4. 활용 가능 데이터", False),
    ("Part II: 알고리즘 설계", True),
    ("  5. 알고리즘 아키텍처: 5-Phase Pipeline", False),
    ("  6. Phase 1: 데이터 수집", False),
    ("  7. Phase 2: 저자 역할 판별", False),
    ("Part III: LLM 자연어 프롬프트 알고리즘 (신규)", True),
    ("  8. LLM 기반 저자 식별 전략", False),
    ("  9. 수학 vs LLM 비교 분석", False),
    ("  10. 권장 전략: 하이브리드 접근법", False),
    ("Part IV: 저자 동일성 검증", True),
    ("  11. 문제 정의", False),
    ("  12. 다중 신호 기반 판별 (5-Signal)", False),
    ("  13. 수학 모델 vs LLM 모델", False),
    ("  14. 동일저자 추적 시나리오", False),
    ("Part V: 매칭 및 업데이트", True),
    ("  15. 5-Level 퍼지 매칭", False),
    ("  16. 최종 매칭 공식", False),
    ("  17. 업데이트 스키마", False),
    ("  18. 산출물", False),
    ("Part VI: 온톨로지 데이터 구조 (Biblo 전략)", True),
    ("  19. Biblo 온톨로지 전략 포인트", False),
    ("  20. 데이터 구조 설계 원칙", False),
    ("  21. 전체 ERD", False),
    ("  22. 핵심 테이블 명세", False),
    ("  23. Biblo <-> KCI 패턴 매핑 총괄", False),
    ("Part VII: 월별 RIMS 덤프 파이프라인", True),
    ("  24. 전체 흐름", False),
    ("  25. 증분 처리 전략", False),
    ("  26. 월별 예상 처리량", False),
    ("Part VIII: 기술 스택 및 실행 계획", True),
    ("  27. 핵심 라이브러리", False),
    ("  28. 모듈 구조", False),
    ("  29. 실행 계획 (5-Sprint)", False),
    ("Part IX: 리스크, KPI, 결론", True),
    ("  30. 리스크 및 대응", False),
    ("  31. 성공 지표 (KPI)", False),
    ("  32. 기대 효과", False),
    ("  33. 온톨로지 성숙도 로드맵", False),
    ("  34. 결론", False),
    ("부록 A: RIMS 데이터 필드 명세", True),
]
for item_text, is_bold in toc_items:
    if is_bold:
        add_para(item_text, size=10, bold=True, space_after=3, color=C_NAVY)
    else:
        add_para(item_text, size=10, space_after=2)

doc.add_page_break()

# ════════════════════════════════════════
#  Executive Summary
# ════════════════════════════════════════
add_heading("Executive Summary", level=1, color=C_NAVY)

add_para("성균관대학교 연구성과관리시스템(RIMS)의 KCI 논문 24,532건에 대해 저자 역할(교신저자, 1저자, 공동1저자, 순위저자)을 정확하게 식별하고, 동명이인 구분 및 동일저자 추적까지 수행하는 자동화 시스템을 구축한다. 이후 결과를 Biblo 온톨로지 전략을 적용한 데이터 구조로 축적하여 연구자 지식 기반을 구축한다.")

add_para("핵심 전략:", bold=True, size=11)
add_bullet("KCI Open API + PDF 1페이지 파싱 이중 경로로 저자 역할 식별")
add_bullet("수학 알고리즘과 LLM 자연어 프롬프트 알고리즘을 병행하여 최적의 식별 정확도 달성")
add_bullet("RIMS 축적 데이터를 활용한 5-Signal 저자 동일성 판별")
add_bullet("5-Phase Pipeline: 수집 → 역할 판별 → 저자 동일성 검증 → RIMS 업데이트 → 온톨로지 구축")
add_bullet("Biblo 온톨로지 전략 적용: ID 기반 엔티티 체계 + 월별 증분 파이프라인")

doc.add_page_break()

# ════════════════════════════════════════
#  Part I: 현황 분석
# ════════════════════════════════════════
add_heading("Part I: 현황 분석", level=1, color=C_NAVY)

# Section 1
add_heading("1. RIMS 데이터 현황", level=2, color=C_BLUE)
make_table(
    ["항목", "수치"],
    [
        ["전체 논문", "136,952건"],
        ["KCI 논문 (대상)", "24,532건"],
        ["저자 참여 레코드", "5,134,679건"],
        ["고유 이름 수", "312,456명"],
        ["고유 연구자 ID (PRTCPNT_ID)", "19,666명"],
        ["PRTCPNT_ID 미보유 레코드", "4,760,125건 (92.7%)"],
    ],
    col_widths=[6, 5]
)

# Section 2
add_heading("2. TPI_DVS_CD (저자 역할 코드) 분포", level=2, color=C_BLUE)
make_table(
    ["코드", "의미", "건수", "정확도"],
    [
        ["2", "1저자", "9,851", "비교적 정확"],
        ["3", "교신저자", "3,797", "일부만 식별"],
        ["4", "공저자", "486,351", "미분류 혼재"],
    ],
    col_widths=[2, 4, 3, 4]
)
add_para("핵심 문제: TPI=4에 교신저자, 2저자, 순위저자가 미분류 상태로 혼재 (486,351건)", bold=True, color=C_RED)

# Section 3
add_heading("3. 동명이인 현황", level=2, color=C_BLUE)
make_table(
    ["항목", "수치"],
    [
        ["동명이인 (같은 이름, 다른 ID)", "13,577명"],
        ["최다 동명이인", "Kim, J. (214명), Lee, J. (210명)"],
        ["소속 변경 이력자 (같은 ID, 2개+ 소속)", "6,167명"],
    ],
    col_widths=[6, 6]
)

# Section 4
add_heading("4. 저자 동일성 판별에 활용 가능한 RIMS 데이터", level=2, color=C_BLUE)
make_table(
    ["데이터", "보유율", "활용 방안"],
    [
        ["논문 제목 (ORG_LANG_PPR_NM)", "100%", "연구 주제 유사도 계산"],
        ["초록 (ABST_CNTN)", "87.0% (KCI 97.4%)", "TF-IDF / 임베딩 기반 유사도"],
        ["학술지명 (SCJNL_NM)", "100%", "연구 분야 추정"],
        ["발행연도 (PBLC_YM)", "100%", "시계열 경력 추적"],
        ["ORCID", "69명", "확정적 식별자 (보유 시)"],
        ["SCOPUS_ID", "1,240명", "확정적 식별자 (보유 시)"],
        ["이메일 (EMAL_ADDR)", "10,437명", "준확정적 식별자"],
        ["공저자 네트워크", "78,509 공저논문", "동명이인 구분의 강력한 신호"],
        ["소속기관 (BLNG_AGC_NM)", "대부분 보유", "시계열 소속 변경 추적"],
    ],
    col_widths=[5, 3, 5]
)

doc.add_page_break()

# ════════════════════════════════════════
#  Part II: 알고리즘 설계
# ════════════════════════════════════════
add_heading("Part II: 알고리즘 설계", level=1, color=C_NAVY)

# Section 5
add_heading("5. 알고리즘 아키텍처: 5-Phase Pipeline", level=2, color=C_BLUE)

add_para("전체 파이프라인 흐름:", bold=True, size=11)
pipeline_text = """Phase 1: 데이터 수집          Phase 2: 저자 역할 판별
  RIMS → KCI ID/DOI 추출        API + PDF 교차 검증
  Track A: KCI API               의사결정 트리 / LLM 프롬프트
  Track C: PDF 1페이지 파싱       역할 분류 (5개 유형)
        │                              │
        ▼                              ▼
Phase 3: 저자 동일성 검증 ◀──── Phase 2 결과
  동명이인 구분 (Disambiguation)
  동일저자 추적 (Entity Resolution)
  수학 점수 모델 / LLM 판단 / 하이브리드
        │
        ▼
Phase 4: 매칭 및 업데이트
  퍼지 매칭 + 동일성 검증 결합
  RIMS 데이터 보정
  최종 산출물 생성
        │
        ▼
Phase 5: 온톨로지 구축 (신규)
  엔티티 ID 부여 (기관/분야/키워드)
  계층 구조 + aliases 매핑
  월별 증분 파이프라인 자동화"""

add_code_block(pipeline_text)

add_para("")
add_para("핵심: 5-Phase 파이프라인. 수학 알고리즘 우선 처리 후, 경계 사례에 대해 LLM 보완. Phase 5에서 온톨로지 ID 체계 구축.", bold=True, color=C_BLUE)

doc.add_page_break()

# Section 6
add_heading("6. Phase 1: 데이터 수집", level=2, color=C_BLUE)

add_heading("6.1 대상 논문 필터링", level=3, color=C_BLUE)
add_bullet("RIMS에서 ID_KCI IS NOT NULL → 24,532건")
add_bullet("DOI + KCI ID 모두 보유: 15,872건")
add_bullet("KCI ID만 보유: 8,660건")

add_heading("6.2 Track A: KCI Open API", level=3, color=C_BLUE)
make_table(
    ["항목", "내용"],
    [
        ["엔드포인트", "open.kci.go.kr/po/openapi/openApiSearch.kci"],
        ["apiCode", "articleDetail"],
        ["출력: author-division", "1 = 주저자, 2 = 공저자"],
        ["출력: author-part", "역할 텍스트 (주저자/공저자/교신저자)"],
        ["추가 출력", "ORCID, 소속"],
        ["제약", "API 키 발급 필요, 일 5,000건 제한"],
    ],
    col_widths=[4, 9]
)

add_heading("6.3 Track C: PDF 1페이지 파싱 (핵심 경로)", level=3, color=C_BLUE)
add_para("처리 단계:", bold=True)
add_bullet("C-1: KCI 웹페이지 접속 → '원문 보러가기' 링크 추출")
add_bullet("C-2: PDF 1페이지 다운로드")
add_bullet("C-3: 텍스트 추출 (PyMuPDF / Tesseract OCR)")
add_bullet("C-4: 저자명 + 역할 기호(* † ‡) 파싱")
add_bullet("C-5: 하단 저자 소개 블록 파싱 (소속, 이메일, ISNI)")

add_para("기호 규칙:", bold=True)
make_table(
    ["기호", "의미"],
    [
        ["*", "교신저자 (corresponding author)"],
        ["**", "공동교신저자"],
        ["†", "공동제1저자 (equal contribution)"],
        ["순서 1번째", "제1저자"],
        ["마지막 순서", "시니어/교신 후보"],
    ],
    col_widths=[3, 9]
)

# Section 7
add_heading("7. Phase 2: 저자 역할 판별", level=2, color=C_BLUE)

add_heading("7.1 수학 알고리즘: 의사결정 트리", level=3, color=C_BLUE)

decision_text = """PDF에서 * 기호?  ──YES──→  교신저자 (순서 1번이면 1저자 겸 교신)
      │ NO
PDF에서 † 기호?  ──YES──→  공동 제1저자
      │ NO
저자 순서 1번?   ──YES──→  1저자
      │ NO
마지막 저자(≥3명)?──YES──→  마지막저자/시니어 (API로 교신 확인)
      │ NO
API '교신' 포함?  ──YES──→  교신저자 (API 기반)
      │ NO
      └──────────→  공저자 + 순서 번호"""

add_code_block(decision_text)

add_heading("7.2 신뢰도 산출", level=3, color=C_BLUE)
make_table(
    ["판별 조건", "신뢰도"],
    [
        ["API + PDF 일치", "1.0"],
        ["PDF 기호만 확인", "0.9"],
        ["API만 확인", "0.8"],
        ["순서 기반 추정", "0.6"],
        ["API ≠ PDF 불일치", "0.4"],
        ["판별 불가", "0.0 → 수동"],
    ],
    col_widths=[7, 3]
)

doc.add_page_break()

# ════════════════════════════════════════
#  Part III: LLM 자연어 프롬프트 알고리즘 (NEW)
# ════════════════════════════════════════
add_heading("Part III: LLM 자연어 프롬프트 알고리즘 (신규)", level=1, color=C_TEAL)

# Section 8
add_heading("8. LLM 기반 저자 식별 전략", level=2, color=C_TEAL)

add_heading("8.1 기본 개념", level=3, color=C_TEAL)
add_para("기존 수학 알고리즘(의사결정 트리, 가중 점수 모델)은 사전 정의된 규칙을 순차적으로 적용하는 방식이다. 반면 LLM 자연어 프롬프트 알고리즘은 PDF/API에서 추출한 텍스트 데이터를 LLM에게 자연어로 맥락을 설명하고, 판단을 요청하는 방식이다.")

add_para('핵심 사상: "수학 공식으로 코딩하지 않아도, 자연어로 판단 기준을 명확히 기술하면 LLM이 복합적 맥락을 이해하여 더 정확한 판단을 내릴 수 있다."', bold=True, color=C_TEAL)

add_heading("8.2 적용 영역별 LLM 프롬프트 설계", level=3, color=C_TEAL)

# Prompt 1: 역할 판별
add_para("영역 1: 저자 역할 판별 (Phase 2 대체/보완)", bold=True, color=C_TEAL, size=11)

prompt1 = """[System Prompt]
당신은 KCI 학술 논문의 저자 역할을 판별하는 전문가입니다.
다음 규칙을 적용하여 각 저자의 역할을 판별하세요:
- * 기호가 붙은 저자 → 교신저자 (corresponding author)
- † 기호가 붙은 저자 → 공동제1저자 (equal contribution)
- 저자 목록의 첫 번째 → 1저자 (first author)
- 저자 목록의 마지막(3명 이상일 때) → 시니어/교신 후보
- 하단 저자정보 블록에 "교신저자" 표기 → 교신저자 확정

[User Prompt]
다음은 KCI 논문 PDF 1페이지에서 추출한 텍스트입니다:

<PDF 텍스트>
{extracted_text}
</PDF 텍스트>

<KCI API 데이터>
{api_response}
</KCI API 데이터>

각 저자에 대해 다음 JSON 형식으로 역할을 판별해주세요:
{
  "authors": [
    {
      "name_kr": "한글명",
      "name_en": "영문명",
      "role": "first_author | corresponding | co_first | last_author | co_author",
      "rank": 순서번호,
      "confidence": 0.0~1.0,
      "evidence": "판별 근거 설명"
    }
  ]
}"""
add_code_block(prompt1, font_size=8)

# Prompt 2: 동명이인 구분
add_para("영역 2: 동명이인 구분 (Phase 3 대체/보완)", bold=True, color=C_TEAL, size=11)

prompt2 = """[System Prompt]
당신은 학술 연구자의 동일성을 판별하는 전문가입니다.
동명이인 구분 시 다음 신호를 종합적으로 고려하세요:
1. 확정 식별자 (ORCID, 이메일) → 일치 시 동일인 확정
2. 공저자 네트워크 → 같은 공저자와 함께 논문 = 동일인 가능성 높음
3. 연구 주제 연속성 → 같은 분야 연구 = 동일인 가능성
4. 소속 변경 타당성 → 시간순으로 자연스러운 경력 진행인지
5. 메타데이터 → 기관명, 이니셜, 부가 식별자 등

[User Prompt]
RIMS에서 "김경수"를 검색하여 다음 3명의 후보를 발견했습니다:

<후보 1: ID=12345>
- 소속: 성균관대 언어AI 전공
- 최근 논문 주제: 자연어처리, 영어교육, 인공지능
- 공저자: 박철수, 이영희, 최민준
- 활동 기간: 2018~2022
</후보>

<후보 2: ID=67890>  ...  <후보 3: ID=11111>  ...

새 논문 정보:
- 제목: "AI 기반 영어 교육 시스템의 효과 분석"
- 소속: 성균관대 영어영문학과 조교수
- 공저자: 박철수, 김민수
- 발행: 2024년

이 논문의 "김경수"는 위 후보 중 누구인지 판별하세요."""
add_code_block(prompt2, font_size=8)

# Prompt 3: OCR 보정
add_para("영역 3: 비정형 텍스트 처리 (OCR 결과, 비표준 형식)", bold=True, color=C_TEAL, size=11)

prompt3 = """[System Prompt]
다음은 OCR로 추출된 학술 논문 PDF의 저자 정보입니다.
OCR 노이즈가 있을 수 있으니, 맥락을 고려하여 저자명과 역할을 판별하세요.

[User Prompt]
OCR 추출 텍스트:
"김 철 수* · 이 영 희** · 박 민 수
*교 신저 자: 성 균관대 학교 소 프트웨 어학 과"

→ OCR 노이즈를 보정하고 저자 역할을 판별해주세요."""
add_code_block(prompt3, font_size=8)

add_heading("8.3 LLM 알고리즘의 구현 아키텍처", level=3, color=C_TEAL)

impl_code = """class LLMAuthorClassifier:
    \"\"\"LLM 기반 저자 역할 판별기\"\"\"

    def __init__(self, model="claude-sonnet-4-5-20250929"):
        self.client = Anthropic()
        self.model = model
        self.system_prompt = ROLE_CLASSIFICATION_PROMPT

    def classify_authors(self, pdf_text, api_data=None):
        \"\"\"PDF 텍스트 + API 데이터로 저자 역할 판별\"\"\"
        user_prompt = self._build_prompt(pdf_text, api_data)
        response = self.client.messages.create(
            model=self.model,
            system=self.system_prompt,
            messages=[{"role": "user", "content": user_prompt}],
            max_tokens=2000
        )
        return self._parse_response(response)


class LLMDisambiguator:
    \"\"\"LLM 기반 동명이인 구분기\"\"\"

    def __init__(self, model="claude-sonnet-4-5-20250929"):
        self.client = Anthropic()
        self.model = model
        self.system_prompt = DISAMBIGUATION_PROMPT

    def disambiguate(self, candidate_name, candidates, new_paper):
        \"\"\"후보 목록에서 새 논문의 저자를 식별\"\"\"
        user_prompt = self._build_context(candidates, new_paper)
        response = self.client.messages.create(
            model=self.model,
            system=self.system_prompt,
            messages=[{"role": "user", "content": user_prompt}],
            max_tokens=1500
        )
        return self._parse_result(response)"""
add_code_block(impl_code, font_size=8)

doc.add_page_break()

# Section 9: 수학 vs LLM 비교 분석
add_heading("9. 수학 알고리즘 vs LLM 알고리즘 비교 분석", level=2, color=C_TEAL)

add_heading("9.1 비교 총괄표", level=3, color=C_TEAL)

# 12-row comparison table
make_table(
    ["평가 항목", "수학 알고리즘", "LLM 프롬프트", "하이브리드 (권장)"],
    [
        ["구현 복잡도", "높음 (모든 케이스 규칙 정의)", "낮음 (프롬프트 작성)", "중간"],
        ["정형 데이터 처리", "★★★ 우수 (규칙 기반)", "★★☆ 양호", "★★★"],
        ["비정형 데이터 처리", "★☆☆ 취약 (OCR 노이즈)", "★★★ 우수 (맥락 이해)", "★★★"],
        ["동명이인 구분", "★★☆ (기계적 합산)", "★★★ (맥락적 종합)", "★★★"],
        ["처리 속도", "★★★ 매우 빠름 (ms)", "★☆☆ 느림 (초 단위)", "★★☆"],
        ["비용", "★★★ 무료 (로컬 연산)", "★☆☆ API 호출 비용", "★★☆"],
        ["재현성", "★★★ 완벽 (결정론적)", "★★☆ 높으나 100% 아님", "★★★ (수학으로 검증)"],
        ["엣지 케이스 대응", "★☆☆ 규칙 외 대응 불가", "★★★ 맥락으로 대응", "★★★"],
        ["설명 가능성", "★★★ 점수 산식 투명", "★★★ 자연어 근거 제공", "★★★"],
        ["유지보수", "★☆☆ 코드 수정 필요", "★★★ 프롬프트 수정만", "★★☆"],
        ["대규모 배치", "★★★ 24K건 1시간 이내", "★☆☆ 24K건 수일 소요", "★★☆"],
        ["감사 추적성", "★★★ 점수+근거 기록", "★★☆ JSON 기록 가능", "★★★"],
    ],
    col_widths=[3.5, 3.5, 3.5, 3.5],
    header_color="00897B"
)

# Section 9.2: 비용 분석
add_heading("9.2 비용 분석", level=3, color=C_TEAL)
make_table(
    ["항목", "수학 알고리즘", "LLM (Claude Sonnet 4.5)"],
    [
        ["24,532건 처리", "서버 비용 ~$0", "API 호출 ~$50~100"],
        ["동명이인 13,577건", "서버 비용 ~$0", "API 호출 ~$30~60"],
        ["월별 증분 (~2,000건)", "~$0", "~$5~10/월"],
        ["연간 운영비", "서버비만", "~$60~120 추가"],
    ],
    col_widths=[5, 4, 5],
    header_color="00897B"
)

# Section 9.3: 적합 영역 분석
add_heading("9.3 적합 영역 분석", level=3, color=C_TEAL)

add_para("수학 알고리즘이 더 적합한 영역:", bold=True, color=C_BLUE, size=11)
add_bullet("정형화된 규칙 적용: PDF * 기호 → 교신저자 (규칙이 명확)", bold_prefix="1. ")
add_bullet("대규모 배치 처리: 24,532건 전체를 빠르게 처리", bold_prefix="2. ")
add_bullet("확정 식별자 매칭: ORCID, 이메일 등 정확한 1:1 매칭", bold_prefix="3. ")
add_bullet("퍼지 이름 매칭: Levenshtein 거리 등 정량적 유사도 계산", bold_prefix="4. ")
add_bullet("공저자 Jaccard 유사도: 집합 연산 기반 수치 계산", bold_prefix="5. ")

add_para("LLM 알고리즘이 더 적합한 영역:", bold=True, color=C_TEAL, size=11)
add_bullet("OCR 노이즈 처리: 깨진 텍스트에서도 맥락으로 저자명/역할 추출", bold_prefix="1. ")
add_bullet("비표준 저자 표기: 학회마다 다른 저자 표기 관행 이해", bold_prefix="2. ")
add_bullet("복합적 동명이인 판별: 5개 신호를 사람처럼 종합적으로 판단", bold_prefix="3. ")
add_bullet('소속 변경 타당성 평가: "연구교수→조교수"의 자연스러움을 맥락으로 판단', bold_prefix="4. ")
add_bullet("애매한 케이스 처리: 수학 모델로 0.4~0.6점 나오는 경계 사례", bold_prefix="5. ")
add_bullet("새로운 패턴 대응: 코드 수정 없이 프롬프트 조정으로 즉시 대응", bold_prefix="6. ")

doc.add_page_break()

# Section 10: 하이브리드 전략
add_heading("10. 권장 전략: 하이브리드 접근법", level=2, color=C_TEAL)

add_heading("10.1 하이브리드 아키텍처", level=3, color=C_TEAL)

hybrid_arch = """┌─────────────────────────────────────────────────────────────┐
│                      입력 데이터                              │
│    PDF 추출 텍스트 + KCI API + RIMS 기존 데이터                │
└──────────────────────────┬──────────────────────────────────┘
                           │
                           ▼
┌─ 1차: 수학 알고리즘 (빠르고 저비용) ───────────────────────┐
│  의사결정 트리 → 역할 판별                                   │
│  확정 ID 매칭 → 동일인 확정                                  │
│  가중 점수 모델 → 동일성 점수 산출                            │
│                                                             │
│  결과: 신뢰도 ≥ 0.7 → 자동 확정 (전체의 ~70%)               │
│        신뢰도 < 0.7 → 2차 판별로 이관 (~30%)                │
└──────────────────────────┬──────────────────────────────────┘
                           │ 신뢰도 < 0.7 케이스
                           ▼
┌─ 2차: LLM 프롬프트 (정확하고 맥락적) ─────────────────────┐
│  자연어 프롬프트로 맥락 전달                                  │
│  LLM이 종합적으로 판단 + 근거 제공                           │
│  OCR 노이즈 보정, 비표준 형식 처리                           │
│                                                             │
│  결과: 신뢰도 ≥ 0.7 → 자동 확정                             │
│        신뢰도 < 0.5 → 수동 검토 이관                        │
└──────────────────────────┬──────────────────────────────────┘
                           │
                           ▼
┌─ 교차 검증 ────────────────────────────────────────────────┐
│  수학 점수와 LLM 판단이 일치 → 최종 신뢰도 상향              │
│  불일치 → 수동 검토 대상으로 플래그                           │
└─────────────────────────────────────────────────────────────┘"""

add_code_block(hybrid_arch, font_size=8)

# Section 10.2: 기대 효과 비교 테이블
add_heading("10.2 하이브리드 전략의 기대 효과", level=3, color=C_TEAL)

make_table(
    ["항목", "수학 단독", "LLM 단독", "하이브리드"],
    [
        ["자동 처리율", "~70%", "~85%", "~88%"],
        ["수동 검토 비율", "~30%", "~15%", "~12%"],
        ["처리 시간 (24K건)", "~1시간", "~50시간", "~5시간"],
        ["API 비용", "$0", "~$100", "~$30 (30%만 LLM)"],
        ["정확도 (정형)", "95%", "93%", "95%"],
        ["정확도 (비정형)", "60%", "90%", "90%"],
        ["정확도 (동명이인)", "80%", "88%", "90%"],
    ],
    col_widths=[4, 3, 3, 3.5],
    header_color="00897B"
)

add_heading("10.3 최종 권장", level=3, color=C_TEAL)
add_para('"수학 알고리즘 우선, LLM 보완" 하이브리드 전략을 권장한다.', bold=True, color=C_TEAL, size=11)

add_para("이유:", bold=True)
add_bullet("비용 효율: 전체의 ~70%는 수학 알고리즘으로 무료 처리, LLM은 ~30%에만 투입", bold_prefix="1. ")
add_bullet("속도: 대규모 배치에서 수학 알고리즘이 압도적으로 빠름", bold_prefix="2. ")
add_bullet("재현성: 수학 알고리즘의 결정론적 결과 + LLM의 감사 추적 결합", bold_prefix="3. ")
add_bullet("정확도 극대화: 수학 모델의 경계 사례를 LLM이 보완하여 전체 정확도 향상", bold_prefix="4. ")
add_bullet("유지보수 용이: 새로운 패턴은 프롬프트 수정으로 즉시 대응 가능", bold_prefix="5. ")

doc.add_page_break()

# ════════════════════════════════════════
#  Part IV: 저자 동일성 검증
# ════════════════════════════════════════
add_heading("Part IV: 저자 동일성 검증", level=1, color=C_NAVY)

# Section 11
add_heading("11. 문제 정의", level=2, color=C_BLUE)

add_para("과제 A: 동명이인 구분 (Name Disambiguation)", bold=True, color=C_PURPLE, size=11)
add_bullet('"Kim, J."라는 이름이 214명의 서로 다른 연구자를 지칭')
add_bullet("PRTCPNT_ID가 없는 레코드가 92.7%")
add_bullet("같은 이름이라도 다른 사람임을 구분해야 함")

add_para("과제 B: 동일저자 추적 (Entity Resolution)", bold=True, color=C_PURPLE, size=11)
add_bullet('같은 연구자가 소속/직위 변경으로 다르게 기술됨')
add_bullet('예: "김경수, 성균관대 언어AI 전공 연구교수" → "김경수, 성균관대 영어영문학과 조교수"')
add_bullet("6,167명이 2개 이상 소속 기관 보유")

# Section 12
add_heading("12. 다중 신호 기반 저자 동일성 판별 (Multi-Signal Identity Resolution)", level=2, color=C_BLUE)

add_para("5개 신호를 종합하여 동일 저자 여부를 판별하는 가중 점수 모델:", bold=True)

p = doc.add_paragraph()
run = p.add_run("Identity Score = w1*ID신호 + w2*네트워크신호 + w3*주제신호 + w4*시계열신호 + w5*메타신호")
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = C_BLUE
run.font.name = 'Consolas'

# Signal 1
add_para("")
add_para("Signal 1: 확정 식별자 (Deterministic Identifiers) — 최우선", bold=True, color=C_BLUE, size=11)
make_table(
    ["식별자", "보유율", "효과"],
    [
        ["PRTCPNT_ID", "7.3%", "같은 ID → 동일인 확정"],
        ["ORCID", "0.4%", "같은 ORCID → 동일인 확정"],
        ["SCOPUS_ID", "6.3%", "같은 SCOPUS → 동일인 확정"],
        ["이메일", "53.1% (ID보유자)", "같은 이메일 → 동일인 확정"],
    ],
    col_widths=[3, 3, 6],
    header_color="1B7A4D"
)
add_para("PRTCPNT_ID가 있으면 동일인 여부를 즉시 확정. 가장 강력한 1차 판별 수단.")

# Signal 2
add_para("Signal 2: 공저자 네트워크 (Co-author Network) — w2 = 0.30", bold=True, color=C_BLUE, size=11)
add_bullet("같은 이름의 두 레코드가 같은 공저자와 함께 논문을 쓴 적이 있으면 → 동일인 가능성 극히 높음")
add_bullet("공저자 Jaccard 유사도: J(A,B) = |CoAuthors_A ∩ CoAuthors_B| / |CoAuthors_A ∪ CoAuthors_B|")
add_bullet("RIMS에 78,509건의 공저 논문 존재 → 강력한 네트워크 신호")

# Signal 3
add_para("Signal 3: 연구 주제 유사도 (Research Topic Similarity) — w3 = 0.25", bold=True, color=C_BLUE, size=11)
add_bullet("RIMS의 초록(ABST_CNTN, 97.4% 보유), 논문 제목, 학술지명을 활용")
add_bullet("방법 1: TF-IDF + 코사인 유사도 — 저자 논문 집합의 centroid 비교")
add_bullet("방법 2: 학술지 분야 매칭 — ISSN 기반 학문 분야 매핑")

# Signal 4
add_para("Signal 4: 시계열 경력 추적 (Temporal Career Tracking) — w4 = 0.20", bold=True, color=C_BLUE, size=11)
add_bullet("RIMS의 발행연도(PBLC_YM)로 연구 활동 타임라인 구성")
add_bullet("소속 변경이 시간순으로 자연스러운지 검증")

# Signal 5
add_para("Signal 5: 메타데이터 매칭 (Metadata Matching) — w5 = 0.10", bold=True, color=C_BLUE, size=11)
add_bullet("이니셜 일치 (Full name에서 추출)")
add_bullet("소속기관 부분 일치 (같은 대학 내 다른 학과)")
add_bullet("KRI_ID, RESEARCHER_ID 등 부가 식별자")

# Section 13
add_heading("13. 동일성 판별 — 수학 모델 vs LLM 모델", level=2, color=C_BLUE)

add_para("수학 모델 (기본):", bold=True, color=C_BLUE, size=11)
math_model = """total = 0.30*s_network + 0.25*s_topic + 0.20*s_temporal + 0.10*s_meta
if total >= 0.7 and (best - second) >= 0.2:
    return "자동 확정"
"""
add_code_block(math_model, font_size=9)

add_para("LLM 모델 (보완):", bold=True, color=C_TEAL, size=11)
add_para("수학 모델로 0.4~0.7 점수 구간의 경계 사례에 대해 LLM에게 후보 정보를 자연어로 제공하고, 종합 판단을 요청. LLM은 수치로 표현하기 어려운 맥락적 판단(소속 변경의 자연스러움, 연구 분야 전환의 합리성 등)에서 수학 모델 대비 우위.")

add_para("판별 기준:", bold=True)
make_table(
    ["조건", "처리"],
    [
        ["점수 >= 0.7 & 1위-2위 차이 >= 0.2", "자동 확정"],
        ["점수 >= 0.5 & 차이 < 0.2", "수동 검토 필요 (동명이인 의심)"],
        ["점수 0.4~0.7 (경계 사례)", "LLM 보완 판별 (하이브리드)"],
        ["점수 < 0.5", "RIMS 미등록 연구자 가능성"],
    ],
    col_widths=[6, 7],
    header_color="6C3EB6"
)

doc.add_page_break()

# Section 14
add_heading("14. 동일저자 추적 시나리오 (예시)", level=2, color=C_BLUE)

add_para('시나리오: "김경수" 소속 변경', bold=True, size=11)
make_table(
    ["연도", "소속 (PDF 추출)", "RIMS 기존 데이터"],
    [
        ["2022", "성균관대 언어AI 전공 연구교수", "PRTCPNT_ID=12345"],
        ["2024", "성균관대 영어영문학과 조교수", "??? (새 논문)"],
    ],
    col_widths=[2, 5, 5]
)

add_para("판별 과정:", bold=True)
add_bullet('"김경수"로 RIMS 검색 → 후보 3명 발견 (ID=12345, 67890, 11111)')
add_bullet("Signal 1: 확정 ID 없음")
add_bullet("Signal 2: ID=12345의 이전 공저자 중 2명이 새 논문에도 등장 → 공저자 네트워크 강한 일치 (+0.3)")
add_bullet("Signal 3: ID=12345의 이전 논문 주제(언어학, AI)와 새 논문 주제(영어교육, AI) 유사 (+0.2)")
add_bullet("Signal 4: 2022→2024 연구교수→조교수 전환은 시계열적으로 타당 (+0.15)")
add_bullet("Signal 5: 동일 대학 (+0.10)")
add_para("종합 점수: 0.75 → ID=12345로 확정", bold=True, color=C_GREEN, size=11)

doc.add_page_break()

# ════════════════════════════════════════
#  Part V: 매칭 및 업데이트
# ════════════════════════════════════════
add_heading("Part V: 매칭 및 업데이트", level=1, color=C_NAVY)

# Section 15
add_heading("15. 5-Level 퍼지 매칭", level=2, color=C_BLUE)
make_table(
    ["Level", "방법", "신뢰도"],
    [
        ["L1", "완전 일치 (정규화 후)", "1.0"],
        ["L2", "성 + 이니셜 매칭", "0.9"],
        ["L3", "한글 ↔ 영문 교차 (로마자 변환)", "0.8"],
        ["L4", "소속 + 부분이름 + ORCID", "0.7"],
        ["L5", "편집거리 (Levenshtein, rapidfuzz)", "0.5~"],
    ],
    col_widths=[2, 6, 2]
)

# Section 16
add_heading("16. 최종 매칭 = 이름 매칭 x 동일성 검증", level=2, color=C_BLUE)
add_para("최종 신뢰도 = Name_Match_Score x Identity_Score", bold=True, color=C_BLUE)
add_bullet("이름 매칭 (Phase 4): 추출된 이름이 RIMS 이름과 얼마나 일치하는가")
add_bullet("동일성 검증 (Phase 3): 같은 이름의 여러 후보 중 실제로 누구인가")

# Section 17
add_heading("17. 업데이트 스키마", level=2, color=C_BLUE)
make_table(
    ["필드명", "타입", "구분", "설명"],
    [
        ["TPI_DVS_CD", "int", "보정", "2=1저자, 3=교신, 4=공저자, 5=공동1저자"],
        ["AUTHOR_ROLE", "varchar", "신규", "first / corresponding / co_first / last / co_author"],
        ["AUTHOR_ORDER", "int", "신규", "논문 내 저자 순서"],
        ["ROLE_CONFIDENCE", "float", "신규", "역할 판별 신뢰도 (0.0~1.0)"],
        ["ROLE_SOURCE", "varchar", "신규", "api / pdf / llm / hybrid / manual"],
        ["MATCH_CONFIDENCE", "float", "신규", "이름 매칭 신뢰도"],
        ["IDENTITY_CONFIDENCE", "float", "신규", "동일성 검증 신뢰도"],
        ["IDENTITY_EVIDENCE", "varchar", "신규", "동일성 판별 근거"],
        ["IDENTITY_METHOD", "varchar", "신규", "deterministic / math / llm / hybrid / manual"],
    ],
    col_widths=[3.5, 1.5, 2, 6]
)

# Section 18
add_heading("18. 산출물", level=2, color=C_BLUE)
add_bullet("보정된 RIMS 저자 데이터", bold_prefix="rims_article_parti_updated.csv: ")
add_bullet("수동 검토 대상 (신뢰도 < 0.7)", bold_prefix="manual_review_list.xlsx: ")
add_bullet("동명이인 판별 결과 상세", bold_prefix="disambiguation_report.xlsx: ")
add_bullet("처리 통계 리포트", bold_prefix="processing_report.html: ")

doc.add_page_break()

# ════════════════════════════════════════
#  Part VI: 온톨로지 데이터 구조
# ════════════════════════════════════════
add_heading("Part VI: 온톨로지 데이터 구조 (Biblo 전략 적용)", level=1, color=C_TEAL)

# Section 19
add_heading("19. Biblo 온톨로지 전략 포인트 추출", level=2, color=C_TEAL)

add_para("Biblo 프로젝트(전략 문서 v3.0, Supabase 스키마, 백서)에서 추출한 핵심 전략 포인트 8가지:", bold=True)

make_table(
    ["#", "전략 포인트", "출처", "KCI 적용성"],
    [
        ["SP-1", "FRBR 2계층 모델 (Work → Manifestation)", "백서, 전략v3", "★★★"],
        ["SP-2", "온톨로지 성숙도 레벨 모델 (L0→L3)", "전략v3", "★★★"],
        ["SP-3", 'ID 기반 의미 검색 ("ID 부여 → 벡터 없이 의미 검색")', "전략v3", "★★★"],
        ["SP-4", "엔티티 계층구조 (parent_id 자기참조)", "스키마, 전략v3", "★★★"],
        ["SP-5", "aliases 컬럼을 통한 텍스트→ID 정규화", "전략v3", "★★★"],
        ["SP-6", "N:M 관계 테이블 (Junction Table) 패턴", "스키마", "★★★"],
        ["SP-7", "entity_relations (엔티티 간 교차 관계)", "스키마", "★★☆"],
        ["SP-8", "5단계 데이터 파이프라인 + SHACL 검증", "백서, 전략v3", "★★☆"],
    ],
    col_widths=[1.5, 6, 3, 2.5],
    header_color="00897B"
)

add_para("SP-1: FRBR 2계층 모델", bold=True, color=C_TEAL, size=11)
add_bullet("Biblo: Work(작품, 314K) → Book/Manifestation(도서, 360K) 1:N 관계")
add_bullet("KCI 적용: Article(논문) → Author Record(저자 출현 레코드) 1:N + Author Identity(식별된 저자) 엔티티")

# 온톨로지 성숙도 레벨 테이블
add_para("SP-2: 온톨로지 성숙도 레벨 모델", bold=True, color=C_TEAL, size=11)
make_table(
    ["레벨", "Biblo 상태", "KCI 현재 상태", "KCI 목표"],
    [
        ["L0", "테이블만 존재", "—", "—"],
        ["L1", "메타데이터 존재", "저자명/소속 텍스트 존재", "현재 수준"],
        ["L1.5", "데이터 있으나 ID 미부여", "알고리즘 역할 분류 완료", "Phase 1~3 후"],
        ["L2", "ID + 계층 + 관계", "저자 ID, 기관/분야 계층", "1차 목표"],
        ["L3", "OWL/SPARQL/LOD", "연구자 그래프, LOD 연계", "장기 목표"],
    ],
    col_widths=[2, 3.5, 4, 3.5],
    header_color="00897B"
)

add_para("SP-3~SP-8 요약:", bold=True, color=C_TEAL, size=11)
add_bullet('"ID를 부여하면 벡터 없이 의미 검색이 된다" → 동명이인 해소의 근본 해법', bold_prefix="SP-3: ")
add_bullet("parent_id 자기참조 → 기관/분야 계층 구조", bold_prefix="SP-4: ")
add_bullet("aliases[] 컬럼 → 저자명/기관명 변형 매핑", bold_prefix="SP-5: ")
add_bullet("N:M Junction 테이블 → 저자-논문-기관-분야 관계", bold_prefix="SP-6: ")
add_bullet("entity_relations → 엔티티 간 교차 관계", bold_prefix="SP-7: ")
add_bullet("5단계 파이프라인 + SHACL 검증 → 월별 증분 처리", bold_prefix="SP-8: ")

doc.add_page_break()

# Section 20
add_heading("20. 데이터 구조 설계 원칙", level=2, color=C_TEAL)
make_table(
    ["원칙", "설명", "출처"],
    [
        ["원본 보존", "RIMS 덤프 원본은 변경하지 않고 별도 보존", 'Biblo "Raw Preservation"'],
        ["2계층 분리", "원본(Raw) ↔ 정제(Refined) 레이어 분리", "SP-1 FRBR"],
        ["ID 중심 설계", "모든 엔티티에 고유 ID 부여", "SP-3"],
        ["계층 구조", "parent_id로 분류 체계 표현", "SP-4"],
        ["별칭 지원", "aliases/name_variants로 텍스트→ID 매핑", "SP-5"],
        ["Junction 관계", "N:M은 반드시 관계 테이블로 표현", "SP-6"],
        ["스냅샷 관리", "월별 덤프를 스냅샷으로 관리, 증분 처리", "월 1회 요건"],
    ],
    col_widths=[3, 6, 3.5],
    header_color="00897B"
)

# Section 21: ERD
add_heading("21. 전체 ERD", level=2, color=C_TEAL)

erd_text = """┌─ RAW LAYER (원본 보존) ──────────────────────────────────────┐
│  rims_dump ──1:N──→ rims_article_raw ──1:N──→ rims_author_raw│
│  (덤프 메타)          (논문 원본)              (저자 원본)      │
└──────────────────────────┬────────────────────┬───────────────┘
                           │ ETL               │ ETL
┌─ REFINED LAYER (정제/온톨로지) ──────────────────────────────┐
│  article ◄──── author_article ────► author_identity          │
│  (논문)   1:N   (저자-논문 관계)  N:1   (식별 저자)           │
│    │              - role, rank, confidence     │              │
│    ├─N:M─→ keyword          ├─N:M─→ institution              │
│    └─N:1─→ journal          ├─N:M─→ research_field           │
│                             └─N:M─→ coauthor_network         │
│  entity_relations (엔티티 간 교차 관계)                       │
│  identification_log (식별 이력/감사)                          │
└──────────────────────────────────────────────────────────────┘"""

add_code_block(erd_text, font_size=8)

doc.add_page_break()

# Section 22: 핵심 테이블 명세
add_heading("22. 핵심 테이블 명세", level=2, color=C_TEAL)

add_para("RAW LAYER", bold=True, color=C_TEAL, size=11)

add_para("rims_dump", bold=True, color=C_BLUE)
make_table(
    ["컬럼", "설명"],
    [
        ["id", "덤프 고유 ID"],
        ["dump_date", "덤프 일자"],
        ["dump_type", "덤프 유형"],
        ["file_article / file_author", "원본 파일 경로"],
        ["article_count / author_count", "레코드 수"],
        ["status", "처리 상태"],
        ["created_at / processed_at", "타임스탬프"],
    ],
    col_widths=[5, 8],
    header_color="00897B"
)

add_para("rims_article_raw", bold=True, color=C_BLUE)
make_table(
    ["컬럼", "설명"],
    [
        ["id", "레코드 ID"],
        ["dump_id (FK)", "덤프 참조"],
        ["rims_paper_id", "RIMS 논문 ID"],
        ["org_lang_ppr_nm / eng_ppr_nm", "논문 제목 (한글/영문)"],
        ["scjnl_nm / pblc_ym / doi", "학술지/발행일/DOI"],
        ["abst_cntn", "초록"],
        ["kci_yn", "KCI 여부"],
        ["raw_json", "원본 JSON"],
    ],
    col_widths=[5, 8],
    header_color="00897B"
)

add_para("rims_author_raw", bold=True, color=C_BLUE)
make_table(
    ["컬럼", "설명"],
    [
        ["id", "레코드 ID"],
        ["dump_id (FK)", "덤프 참조"],
        ["rims_paper_id", "RIMS 논문 ID"],
        ["prtcpnt_id", "참여자 ID"],
        ["author_name_kr / author_name_en", "저자명 (한글/영문)"],
        ["tpi_dvs_cd", "저자 역할 코드"],
        ["org_nm / dept_nm", "소속 기관/학과"],
        ["raw_json", "원본 JSON"],
    ],
    col_widths=[5, 8],
    header_color="00897B"
)

add_para("REFINED LAYER", bold=True, color=C_TEAL, size=11)

add_para("article", bold=True, color=C_BLUE)
make_table(
    ["컬럼", "설명"],
    [
        ["id", "논문 고유 ID"],
        ["rims_paper_id (UNIQUE)", "RIMS 논문 ID"],
        ["title_kr / title_en", "논문 제목"],
        ["journal_id (FK)", "학술지 참조"],
        ["publish_date / doi / abstract", "발행일/DOI/초록"],
        ["is_kci / author_count", "KCI 여부, 저자 수"],
        ["first_dump_id / last_dump_id", "최초/최종 덤프 ID"],
    ],
    col_widths=[5, 8],
    header_color="00897B"
)

add_para("author_identity (핵심)", bold=True, color=C_BLUE)
make_table(
    ["컬럼", "설명"],
    [
        ["id", "저자 고유 ID"],
        ["canonical_name_kr / canonical_name_en", "대표명 (한글/영문)"],
        ["name_variants[]", "이름 변형 목록"],
        ["prtcpnt_ids[]", "연결된 RIMS ID 목록"],
        ["primary_institution_id (FK)", "주 소속기관"],
        ["primary_field_id (FK)", "주 연구분야"],
        ["orcid", "ORCID"],
        ["article_count / confidence_score", "논문 수, 신뢰도"],
        ["status", "active / merged / deprecated"],
    ],
    col_widths=[5, 8],
    header_color="00897B"
)

doc.add_page_break()

add_para("온톨로지 엔티티 (계층구조)", bold=True, color=C_TEAL, size=11)

add_para("institution", bold=True, color=C_BLUE)
make_table(
    ["컬럼", "설명"],
    [
        ["id / name / aliases[]", "ID, 기관명, 별칭 목록"],
        ["parent_id (self FK)", "상위 기관 참조 (계층)"],
        ["institution_type", "university / college / department"],
        ["country", "국가"],
    ],
    col_widths=[5, 8],
    header_color="00897B"
)

add_para("계층 구조 예시:", bold=True, size=10)
inst_example = """성균관대학교 (id=1, parent_id=NULL)
  ├── 공과대학 (id=10, parent_id=1)
  │   └── 소프트웨어학과 (id=101, parent_id=10)
  └── 자연과학대학 (id=20, parent_id=1)"""
add_code_block(inst_example, font_size=9)

add_para("research_field", bold=True, color=C_BLUE)
make_table(
    ["컬럼", "설명"],
    [
        ["id / name / aliases[]", "ID, 분야명, 별칭 목록"],
        ["parent_id (self FK)", "상위 분야 참조 (계층)"],
        ["field_code / level", "분야 코드, 계층 레벨"],
    ],
    col_widths=[5, 8],
    header_color="00897B"
)

add_para("keyword", bold=True, color=C_BLUE)
make_table(
    ["컬럼", "설명"],
    [
        ["id / label / aliases[]", "ID, 키워드, 별칭 목록"],
        ["parent_id (self FK)", "상위 키워드 참조"],
    ],
    col_widths=[5, 8],
    header_color="00897B"
)

add_para("관계 테이블 (N:M Junction)", bold=True, color=C_TEAL, size=11)

make_table(
    ["테이블", "컬럼", "설명"],
    [
        ["author_article", "author_id, article_id, role, author_rank, confidence, source, raw_author_id", "저자-논문 관계"],
        ["author_institution", "author_id, institution_id, period_start/end, is_primary", "저자-기관 (시간)"],
        ["author_field", "author_id, field_id, weight, article_count", "저자-분야"],
        ["article_keyword", "article_id, keyword_id, relevance", "논문-키워드"],
        ["coauthor_network", "author_id_1, author_id_2, collaboration_count, first/last", "공저자 네트워크"],
        ["entity_relations", "source_type/id, relation_type, target_type/id, metadata", "엔티티 교차 관계"],
    ],
    col_widths=[3, 7.5, 3],
    header_color="00897B"
)

add_para("이력 테이블", bold=True, color=C_TEAL, size=11)
make_table(
    ["테이블", "핵심 컬럼", "설명"],
    [
        ["identification_log", "dump_id, author_identity_id, raw_author_id", "식별 이력/감사"],
        ["", "action, algorithm_phase, confidence", "식별 행위 기록"],
        ["", "signals_used(jsonb), previous_state(jsonb)", "신호 + 이전 상태"],
    ],
    col_widths=[3.5, 5.5, 4],
    header_color="00897B"
)

doc.add_page_break()

# Section 23
add_heading("23. Biblo <-> KCI 패턴 매핑 총괄", level=2, color=C_TEAL)
make_table(
    ["Biblo 패턴", "Biblo 적용", "KCI 적용", "비고"],
    [
        ["FRBR Work", "작품 (314K)", "article (논문)", "추상적 지적 단위"],
        ["FRBR Manifestation", "book (360K)", "rims_article_raw (원본)", "물리적 발현"],
        ["—", "—", "author_identity (저자)", "고유 엔티티"],
        ["genre", "장르 (13K)", "research_field (연구분야)", "분류 체계"],
        ["concept", "개념 (374K)", "keyword (키워드)", "주제어"],
        ["place", "장소 (80K)", "institution (기관)", "소속 기관"],
        ["work_genres", "작품-장르 (630K)", "author_field (저자-분야)", "N:M 관계"],
        ["work_concepts", "작품-개념 (1.5M)", "article_keyword (논문-키워드)", "N:M 관계"],
        ["work_places", "작품-장소 (457K)", "author_institution (저자-기관)", "N:M + 시간"],
        ["entity_relations", "교차 (326)", "entity_relations", "동일 패턴"],
    ],
    col_widths=[3, 3.5, 4, 3],
    header_color="00897B"
)

doc.add_page_break()

# ════════════════════════════════════════
#  Part VII: 월별 RIMS 덤프 처리 파이프라인
# ════════════════════════════════════════
add_heading("Part VII: 월별 RIMS 덤프 처리 파이프라인", level=1, color=C_TEAL)

# Section 24
add_heading("24. 전체 흐름", level=2, color=C_TEAL)

pipeline_monthly = """매월 1일: RIMS 덤프 수신
     ▼
STEP 1: 덤프 적재
  rims_dump 레코드 생성 → rims_article_raw/rims_author_raw 벌크 INSERT
     ▼
STEP 2: 증분 감지 (Delta Detection)
  신규/변경/삭제 논문 식별 → article INSERT/UPDATE
     ▼
STEP 3: 저자식별 알고리즘 (Phase 1~4)
  수집 → 역할 판별 → 동일성 검증 → 매칭/업데이트
  수학 알고리즘 1차 → LLM 보완 2차 (하이브리드)
     ▼
STEP 4: 온톨로지 갱신
  institution/research_field/keyword 매핑 + coauthor_network 갱신
     ▼
STEP 5: 품질 검증 (SHACL)
  필수 필드·관계 무결성·중복 검사 → 통계 리포트"""

add_code_block(pipeline_monthly)

# Section 25
add_heading("25. 증분 처리 전략", level=2, color=C_TEAL)
make_table(
    ["구분", "감지 방법", "처리"],
    [
        ["신규 논문", "rims_paper_id NOT IN article", "INSERT + 전체 알고리즘"],
        ["변경 논문", "raw_json diff 비교", "UPDATE + 변경 필드만 재처리"],
        ["삭제 논문", "이전 덤프에 있고 현재 없음", "soft delete"],
        ["신규 저자", "새 논문의 저자 레코드", "하이브리드 식별"],
        ["기존 저자 새 논문", "author_identity 존재 + 신규 article", "author_article INSERT"],
    ],
    col_widths=[3, 5, 5],
    header_color="00897B"
)

# Section 26
add_heading("26. 월별 예상 처리량", level=2, color=C_TEAL)
make_table(
    ["항목", "월 예상", "누적 (1년)"],
    [
        ["신규 논문", "~2,000건", "~24,000건"],
        ["신규 저자 레코드", "~7,400건", "~88,800건"],
        ["수학 알고리즘 처리", "~5,200건 (70%)", "—"],
        ["LLM 보완 처리", "~2,200건 (30%)", "—"],
        ["알고리즘 처리 시간", "~2~4시간", "—"],
    ],
    col_widths=[5, 4, 4],
    header_color="00897B"
)

doc.add_page_break()

# ════════════════════════════════════════
#  Part VIII: 기술 스택 및 실행 계획
# ════════════════════════════════════════
add_heading("Part VIII: 기술 스택 및 실행 계획", level=1, color=C_NAVY)

# Section 27
add_heading("27. 핵심 라이브러리", level=2, color=C_BLUE)
make_table(
    ["라이브러리", "용도", "카테고리"],
    [
        ["pandas", "RIMS CSV 데이터 처리", "데이터"],
        ["requests / lxml", "KCI API/웹 접근", "수집"],
        ["PyMuPDF (fitz)", "PDF 텍스트 추출", "추출"],
        ["pytesseract + Pillow", "OCR (이미지 PDF)", "추출"],
        ["rapidfuzz", "퍼지 문자열 매칭", "매칭"],
        ["korean-romanizer", "한글 로마자 변환", "매칭"],
        ["scikit-learn", "TF-IDF + 코사인 유사도", "동일성"],
        ["networkx", "공저자 네트워크 분석", "동일성"],
        ["anthropic", "Claude API (LLM 프롬프트)", "LLM 판별"],
        ["tqdm", "진행률 표시", "유틸"],
    ],
    col_widths=[4, 5, 2.5]
)

# Section 28
add_heading("28. 모듈 구조", level=2, color=C_BLUE)

modules_text = """author_classify/
├── config.py                # 설정
├── main.py                  # 파이프라인 오케스트레이터
├── data/
│   ├── rims_loader.py       # RIMS CSV 로드
│   └── output_writer.py     # 결과 출력
├── collectors/
│   ├── kci_api.py           # KCI Open API (Track A)
│   ├── kci_scraper.py       # KCI 웹페이지 스크래핑
│   └── pdf_downloader.py    # PDF 다운로드
├── extractors/
│   ├── api_parser.py        # API XML 파싱
│   ├── pdf_extractor.py     # PDF 텍스트 추출
│   └── author_parser.py     # 저자명/역할 파싱
├── matchers/
│   ├── name_matcher.py      # 5-Level 퍼지 매칭
│   ├── role_classifier.py   # 의사결정 트리
│   └── cross_validator.py   # 교차 검증
├── identity/
│   ├── disambiguator.py     # 동명이인 구분
│   ├── entity_resolver.py   # 동일저자 추적
│   ├── coauthor_network.py  # 공저자 네트워크 분석
│   ├── topic_similarity.py  # 연구 주제 유사도
│   └── temporal_tracker.py  # 시계열 경력 추적
├── llm/                        # 신규
│   ├── prompt_templates.py    # 프롬프트 템플릿
│   ├── llm_classifier.py     # LLM 역할 판별
│   ├── llm_disambiguator.py  # LLM 동명이인 구분
│   └── hybrid_engine.py      # 하이브리드 오케스트레이터
├── ontology/                   # 신규
│   ├── entity_manager.py      # 엔티티 CRUD
│   ├── hierarchy_builder.py   # 계층 구조 빌더
│   ├── alias_mapper.py        # aliases 매핑
│   └── incremental_sync.py    # 월별 증분 동기화
└── utils/
    ├── text_normalizer.py   # 텍스트 정규화
    ├── korean_utils.py      # 한글 처리
    └── rate_limiter.py      # API 속도 제한"""

add_code_block(modules_text, font_size=8)

doc.add_page_break()

# Section 29
add_heading("29. 실행 계획 (5-Sprint)", level=2, color=C_BLUE)

add_heading("Sprint 1: 기반 구축 + PDF 경로", level=3, color=C_BLUE)
add_para("(API 키 없이 진행 가능)", color=C_GRAY, size=9)
add_bullet("RIMS 데이터 로더 구현")
add_bullet("PDF URL 추출 + 텍스트 추출기")
add_bullet("저자명/역할 기호 파서")
add_bullet("10건 샘플 테스트 → 파싱 정확도 검증")

add_heading("Sprint 2: 매칭 + 역할 판별", level=3, color=C_BLUE)
add_bullet("퍼지 매칭 (5-Level)")
add_bullet("역할 의사결정 트리")
add_bullet("100건 샘플 테스트 → 정확도 측정")

add_heading("Sprint 3: 저자 동일성 검증", level=3, color=C_PURPLE)
add_bullet("공저자 네트워크 + TF-IDF 엔진")
add_bullet("시계열 경력 추적")
add_bullet("5-Signal 종합 모델")
add_bullet("동명이인 100건 테스트")

add_heading("Sprint 4: LLM 통합 + 하이브리드 (신규)", level=3, color=C_TEAL)
add_bullet("LLM 프롬프트 설계 및 테스트")
add_bullet("하이브리드 오케스트레이터 구현")
add_bullet("수학 vs LLM vs 하이브리드 정확도 비교 실험")
add_bullet("최적 임계값 튜닝 (0.7 기준점 조정)")

add_heading("Sprint 5: KCI API 통합 + 대규모 실행 + 온톨로지", level=3, color=C_TEAL)
add_bullet("KCI API 키 연동 + 교차 검증")
add_bullet("전체 24,532건 배치 실행")
add_bullet("온톨로지 엔티티 초기 구축")
add_bullet("월별 증분 파이프라인 자동화")

doc.add_page_break()

# ════════════════════════════════════════
#  Part IX: 리스크, KPI, 결론
# ════════════════════════════════════════
add_heading("Part IX: 리스크, KPI, 결론", level=1, color=C_NAVY)

# Section 30
add_heading("30. 리스크 및 대응", level=2, color=C_BLUE)
make_table(
    ["이슈", "빈도", "대응"],
    [
        ["API 키 미발급", "—", "PDF 경로만으로 우선 진행"],
        ["PDF 접근 불가", "~20%", "DOI 리다이렉트 → 순서 기반 추정"],
        ["이미지 PDF", "~10%", "Tesseract OCR + LLM OCR 보정"],
        ["저자 기호 비표준", "~5%", "LLM 맥락 기반 파싱 (하이브리드)"],
        ["한영 저자명 불일치", "~15%", "5-Level 퍼지 매칭"],
        ["동명이인", "~4.3%", "5-Signal + LLM 하이브리드"],
        ["소속 변경자", "~31%", "시계열 추적 + LLM 판단"],
        ["PRTCPNT_ID 미보유", "92.7%", "하이브리드 추정"],
        ["LLM API 장애", "드묾", "수학 알고리즘 단독 Fallback"],
    ],
    col_widths=[4, 2, 7],
    header_color="C0392B"
)

add_para("핵심 완화 전략:", bold=True, size=11)
add_bullet("이중 경로 (Dual Track): API 실패 시 PDF 독립 동작")
add_bullet("하이브리드 Fallback: LLM 장애 시 수학 알고리즘 단독 동작")
add_bullet("점진적 배치: 10건 → 100건 → 전체 순차 확대")
add_bullet("수동 검토 분리: 신뢰도 < 0.7 자동 플래그")
add_bullet("Fallback 체인: API → PDF → DOI → OCR → LLM → 수동")

# Section 31
add_heading("31. 성공 지표 (KPI)", level=2, color=C_BLUE)
make_table(
    ["지표", "목표"],
    [
        ["교신저자 식별률", "> 85%"],
        ["1저자 검증률", "> 95%"],
        ["저자명 매칭 정확도", "> 90%"],
        ["동명이인 정확 구분률", "> 85% (하이브리드)"],
        ["동일저자 추적 정확도", "> 88% (하이브리드)"],
        ["수동 검토 비율", "< 12% (하이브리드)"],
        ["전체 처리 시간", "< 72시간 (배치)"],
    ],
    col_widths=[6, 3]
)

# Section 32
add_heading("32. 기대 효과", level=2, color=C_BLUE)
make_table(
    ["항목", "AS-IS", "TO-BE"],
    [
        ["교신저자 데이터", "3,797건", "20,000건+"],
        ["저자 역할 분류", "3단계", "5단계"],
        ["동명이인 처리", "수동/미처리", "자동 판별 (하이브리드)"],
        ["소속 변경 추적", "미지원", "시계열 자동 추적"],
        ["데이터 정확도", "미검증", "신뢰도 기반 품질 관리"],
        ["저자 온톨로지", "없음", "ID 기반 엔티티 체계 (L2)"],
        ["월별 운영", "수동", "자동 증분 파이프라인"],
    ],
    col_widths=[4, 4, 4],
    header_color="1B7A4D"
)

doc.add_page_break()

# Section 33
add_heading("33. 온톨로지 성숙도 로드맵", level=2, color=C_TEAL)

add_para("Phase A (L0→L1): 기반 구축 — 1~2개월", bold=True, color=C_TEAL, size=11)
add_bullet("RAW LAYER 테이블 생성, 첫 RIMS 덤프 적재, 증분 감지 로직")

add_para("Phase B (L1→L1.5): 저자식별 알고리즘 적용 — 2~3개월", bold=True, color=C_TEAL, size=11)
add_bullet("하이브리드 알고리즘 실행, author_article 관계 구축, identification_log 기록")

add_para("Phase C (L1.5→L2): 온톨로지 ID 체계 구축 — 2~3개월", bold=True, color=C_TEAL, size=11)
add_bullet("institution/research_field/keyword 엔티티 구축 (계층 + aliases)")
add_bullet("N:M 관계 테이블 + coauthor_network + entity_relations 완성")
add_bullet('"ID 기반 의미 검색" 달성', bold_prefix="목표: ", color=C_GREEN)

add_para("Phase D (L2→L3): 고도화 — 장기", bold=True, color=C_TEAL, size=11)
add_bullet("벡터 임베딩, OWL 온톨로지, SPARQL, 외부 LOD 연계")

# Section 34
add_heading("34. 결론", level=2, color=C_NAVY)

add_para("본 프로젝트는 KCI 논문 저자 식별이라는 핵심 과제를 수학 알고리즘과 LLM 자연어 프롬프트의 하이브리드 접근법으로 해결하고, 그 결과를 Biblo 온톨로지 전략을 적용한 체계적인 데이터 구조에 축적한다.")

add_para("핵심 결론:", bold=True, size=11, color=C_NAVY)
add_bullet("하이브리드 전략: 수학 알고리즘(빠르고 저비용)으로 70%를 처리하고, LLM(정확하고 맥락적)으로 30%를 보완 → 비용 대비 최고 정확도", bold_prefix="1. ")
add_bullet('온톨로지 L2 목표: ID 기반 엔티티 체계를 구축하여 "SQL만으로 연구자 프로필 검색" 달성', bold_prefix="2. ")
add_bullet("월별 자동화: RIMS 덤프 수신부터 온톨로지 갱신까지 전 과정 자동화", bold_prefix="3. ")
add_bullet("감사 추적: identification_log로 모든 식별 결정의 근거와 방법 기록", bold_prefix="4. ")

doc.add_page_break()

# ════════════════════════════════════════
#  부록 A
# ════════════════════════════════════════
add_heading("부록 A: RIMS 데이터 필드 명세", level=1, color=C_NAVY)

add_heading("rims_article_data.csv (논문 메타데이터)", level=2, color=C_BLUE)
make_table(
    ["필드", "활용"],
    [
        ["ARTICLE_ID", "내부 조인 키"],
        ["ORG_LANG_PPR_NM", "논문 제목 → 주제 유사도"],
        ["ABST_CNTN", "초록 → TF-IDF / LLM 유사도"],
        ["SCJNL_NM", "학술지명 → 분야 추정"],
        ["PBLC_YM", "발행연도 → 시계열 추적"],
        ["DOI / ID_KCI", "외부 데이터 접근 키"],
    ],
    col_widths=[5, 7]
)

add_heading("rims_article_parti_data.csv (저자 참여정보)", level=2, color=C_BLUE)
make_table(
    ["필드", "활용"],
    [
        ["PRTCPNT_ID", "확정 식별자"],
        ["PRTCPNT_NM / PRTCPNT_FULL_NM", "이름 매칭 대상"],
        ["TPI_DVS_CD", "보정 대상"],
        ["BLNG_AGC_NM", "소속 매칭 + 시계열 추적"],
        ["ORCID_ID / SCOPUS_ID", "확정 식별자"],
        ["EMAL_ADDR", "준확정 식별자"],
        ["ARTICLE_ID", "공저자 네트워크 구축"],
    ],
    col_widths=[5, 7]
)

add_para("")
add_para("본 보고서는 RIMS 데이터 품질 고도화 프로젝트의 통합 기술 전략 문서입니다.", color=C_GRAY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("저자 식별 알고리즘 + LLM 비교 분석 + Biblo 온톨로지 전략 적용을 포괄합니다.", color=C_GRAY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("성균관대학교 | 2026.02", color=C_GRAY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)


# ════════════════════════════════════════
#  Save
# ════════════════════════════════════════
output_path = "/Users/ks.kim/Documents/kks/author_classify/KCI_저자식별_통합_전략보고서_FINAL.docx"
doc.save(output_path)
print(f"DOCX 생성 완료: {output_path}")
print(f"파일 크기: {os.path.getsize(output_path):,} bytes")
